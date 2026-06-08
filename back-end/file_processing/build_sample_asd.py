"""Generate a deterministic fictional ASD .docx exercising every parser path.

Run:
    python back-end/file_processing/build_sample_asd.py
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from file_processing.parsing import (
    deployment_details_headers,
    ec2_table_headers,
    requirement_headers,
    servers_table_headers,
    solution_headers,
    summary_headers,
)

OUTPUT = Path(__file__).resolve().parent / "data" / "sample_asd.docx"


SUMMARY_CONTENT = {
    "Introduction": (
        "Project Aurora is a fictional inventory analytics platform supporting "
        "regional warehouses. This document describes the proposed cloud architecture "
        "and operational requirements for review."
    ),
    "Key Functionalities/Capabilities": (
        "Real-time stock telemetry ingestion, batch replenishment forecasting, "
        "and an operator portal exposing dashboards and ad-hoc reporting."
    ),
    "Assumptions/Constraints/Recommendations": (
        "Assumes a stable VPC peering arrangement with on-prem ERP, AWS workloads "
        "in us-east-1 with DR in us-west-2, and operator SSO via the existing IdP."
    ),
}

REQUIREMENTS_CONTENT = {
    "User/Usage Requirements": "Up to 300 concurrent operators across two regions.",
    "Interface Requirements": "REST and gRPC interfaces over mTLS only.",
    "Security Requirements": "Customer-managed keys, audit logs retained 365 days.",
    "Network Requirements": "Private subnets, VPC peering, no public ingress.",
    "Software Requirements": "Linux 22.04 LTS, container runtime, JDK 21.",
    "Performance Requirements": "P95 inbound API latency under 200 ms.",
    "Supportability Requirements": "Centralised log aggregation, runbook for each alert.",
    "Storage Requirements": "Encrypted block storage; 5 TiB OLTP, 20 TiB analytics.",
    "Database Requirements": "Managed Postgres with cross-region read replicas.",
    "Disaster Recovery Requirements": "RPO 15 min, RTO 60 min; documented failover.",
    "Compliance Requirements": "SOC2 controls, data residency in North America.",
    "Licensing Requirements": "Open-source preferred; commercial licences inventoried.",
}

SOLUTION_CONTENT = {
    "Proposed New Architecture": (
        "Three-tier microservice topology behind a regional load balancer; "
        "event ingestion via a managed Kafka cluster; analytics via a managed "
        "lakehouse pattern."
    ),
    "Pre-Production Architecture": (
        "Scaled-down clone of production in an isolated VPC; nightly data refresh "
        "with synthetic operator traffic for canary validation."
    ),
    "Production/DR Architecture": (
        "Active-passive across two regions; database streaming replication; "
        "automated DR drill quarterly."
    ),
}

EC2_ROWS = [
    {
        "Environment": "Production",
        "Account Type": "Workload",
        "Network Zone": "Private-App",
        "AWS Region": "us-east-1",
        "AZ": "us-east-1a",
        "OS": "Ubuntu 22.04",
        "Instance Type CPU/RAM": "m6i.2xlarge (8/32)",
        "Count": "6",
        "Storage Type": "gp3",
        "Storage Volume Size": "200 GiB",
        "Domain/DNS": "aurora.internal",
        "Data Residency Restrictions": "USA",
        "Data Classification": "Internal",
        "Server Role": "API",
        "On/Off Scheduling": "24x7",
    },
    {
        "Environment": "Pre-Prod",
        "Account Type": "Workload",
        "Network Zone": "Private-App",
        "AWS Region": "us-east-1",
        "AZ": "us-east-1b",
        "OS": "Ubuntu 22.04",
        "Instance Type CPU/RAM": "m6i.large (2/8)",
        "Count": "2",
        "Storage Type": "gp3",
        "Storage Volume Size": "100 GiB",
        "Domain/DNS": "aurora-pp.internal",
        "Data Residency Restrictions": "USA",
        "Data Classification": "Internal",
        "Server Role": "API",
        "On/Off Scheduling": "Business hours",
    },
]

SERVER_ROWS = [
    {
        "Environment/Location": "DC-East",
        "Server Type": "Physical",
        "OS": "Windows Server 2022",
        "Network Zone": "Restricted",
        "CPU Cores": "16",
        "RAM": "64 GiB",
        "Non-OS SAN Storage": "2 TiB",
        "Count": "2",
        "Domain/DNS": "aurora.corp",
        "Data Residency Restrictions": "USA",
        "Data Classification": "Confidential",
        "Server Role": "ERP Bridge",
    },
    {
        "Environment/Location": "DC-West",
        "Server Type": "Virtual",
        "OS": "Ubuntu 22.04",
        "Network Zone": "Restricted",
        "CPU Cores": "8",
        "RAM": "32 GiB",
        "Non-OS SAN Storage": "1 TiB",
        "Count": "2",
        "Domain/DNS": "aurora.corp",
        "Data Residency Restrictions": "USA",
        "Data Classification": "Confidential",
        "Server Role": "Reporting",
    },
]

DEPLOYMENT_ROWS = [
    {
        "Hosted Location": "AWS us-east-1 (primary), us-west-2 (DR)",
        "Countries/Regions Serviced": "US, Canada",
        "Business Unit(s)": "Supply Chain Operations",
    },
]


def _add_paragraph(doc, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(text)
    if style:
        p.style = doc.styles[style]


def _add_table(doc, headers: list[str], rows: list[dict]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for c, h in enumerate(headers):
        table.cell(0, c).text = h
    for r, row in enumerate(rows, start=1):
        for c, h in enumerate(headers):
            table.cell(r, c).text = str(row.get(h, ""))


def build() -> Path:
    doc = Document()
    # Title page so first "Summary" mention is the table-of-contents-like
    # reference; second mention is the real section. The parser only acts on
    # the second occurrence.
    _add_paragraph(doc, "Project Aurora — Architecture Solution Design", "Title")
    _add_paragraph(doc, "Sections: Summary, Solution Requirements, Proposed Solution")

    # ── Summary ───────────────────────────────────────────────
    _add_paragraph(doc, "Summary", "Heading 1")
    for h in summary_headers:
        _add_paragraph(doc, h, "Heading 2")
        _add_paragraph(doc, SUMMARY_CONTENT[h])

    # ── Solution Requirements ─────────────────────────────────
    _add_paragraph(doc, "Solution Requirements", "Heading 1")
    for h in requirement_headers:
        _add_paragraph(doc, h, "Heading 2")
        _add_paragraph(doc, REQUIREMENTS_CONTENT[h])

    # End-marker for the Solution Requirements section
    _add_paragraph(doc, "Affinity/Anti-Affinity Requirements", "Heading 1")
    _add_paragraph(doc, "Workloads must be split across two AZs at minimum.")

    # ── Proposed Solution ────────────────────────────────────
    _add_paragraph(doc, "Proposed Solution", "Heading 1")
    for h in solution_headers:
        _add_paragraph(doc, h, "Heading 2")
        _add_paragraph(doc, SOLUTION_CONTENT[h])

    # ── EC2 Sizing/Specifications + table ─────────────────────
    _add_paragraph(doc,
                   "EC2 Sizing/Specifications (Guidance on OS Volumes & MS Office Support)",
                   "Heading 1")
    _add_table(doc, ec2_table_headers, EC2_ROWS)

    # ── On-Prem Servers Sizing/Specifications + table ────────
    _add_paragraph(doc, "On-Prem Servers Sizing/Specifications", "Heading 1")
    _add_table(doc, servers_table_headers, SERVER_ROWS)

    # End-marker for the servers table
    _add_paragraph(doc, "Proposed Server Details", "Heading 1")
    _add_paragraph(doc, "Detailed BOM in appendix.")

    # ── Hosted Location + table ──────────────────────────────
    _add_paragraph(doc, "Hosted Location", "Heading 1")
    _add_table(doc, deployment_details_headers, DEPLOYMENT_ROWS)

    # End-marker
    _add_paragraph(doc, "Miscellaneous Information", "Heading 1")
    _add_paragraph(doc, "End of document.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    p = build()
    print(f"Wrote {p}")
    sys.exit(0)
