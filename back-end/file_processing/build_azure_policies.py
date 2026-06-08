"""Generate ``azure_policies.docx`` - default Azure infra policies (15 sections).

Each section header is UPPERCASE, numbered (``N. TITLE``), so the
header-detection rule shared by ``extract_policies`` (PDF) and
``extract_policies_docx`` (DOCX) picks them up.

All content is originally written, Well-Architected-style guidance — no
copyrighted text. Run:

    python back-end/file_processing/build_azure_policies.py
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


OUTPUT = Path(__file__).resolve().parent / "data" / "azure_policies.docx"


SECTIONS: list[tuple[str, str]] = [
    (
        "1. IDENTITY AND ACCESS MANAGEMENT",
        (
            "Identity is the new perimeter for Azure workloads. All human and "
            "workload identities must be issued from a single Microsoft Entra ID "
            "tenant; multi-tenant workloads should rely on application registrations "
            "with limited consent rather than shared credentials. Privileged access "
            "is granted just-in-time through Privileged Identity Management with "
            "time-bound activation, approval workflows, and access reviews on a "
            "quarterly cadence. Conditional Access policies must enforce "
            "phishing-resistant MFA for all administrators and for any user "
            "accessing production data planes; legacy authentication protocols "
            "are blocked at the tenant level. Workload identities favour managed "
            "identities — system-assigned where the lifecycle matches the resource, "
            "user-assigned where the identity is shared across resources or "
            "deployment slots. Service principals with client secrets are only "
            "permitted as a documented exception and must rotate credentials at "
            "least every 90 days through automation. Role assignments follow the "
            "principle of least privilege: prefer built-in roles, scope to the "
            "smallest resource boundary that satisfies the requirement, and use "
            "Azure ABAC conditions where the platform supports them. Break-glass "
            "accounts are excluded from Conditional Access only for the specific "
            "controls that would lock them out, monitored continuously, and stored "
            "in a sealed offline process. Guest access is governed through entitlement "
            "management with sponsor approval, automatic expiry, and access reviews. "
            "Federation with on-premises Active Directory uses Microsoft Entra "
            "Connect with password hash sync as the default, falling back to "
            "pass-through authentication only where regulatory constraints "
            "prohibit hash synchronisation."
        ),
    ),
    (
        "2. NETWORK SECURITY AND SEGMENTATION",
        (
            "Azure network design starts from a hub-and-spoke topology with an "
            "Azure Virtual WAN or a regional hub virtual network providing shared "
            "services: Azure Firewall, ExpressRoute or VPN gateways, DNS resolvers, "
            "and bastion hosts. Workload spokes are peered to the hub and never to "
            "each other, forcing east-west traffic through inspection. Each spoke "
            "is subdivided into purpose-built subnets (web tier, application tier, "
            "data tier, management) with Network Security Groups applying default-"
            "deny rules and explicit allow lists keyed to Application Security "
            "Groups rather than IP ranges. Public ingress is concentrated on Azure "
            "Front Door or Application Gateway with the integrated Web Application "
            "Firewall in prevention mode using the latest OWASP and Microsoft-managed "
            "rule sets. Outbound internet egress flows through Azure Firewall with "
            "FQDN-based rules; direct workload-to-internet routes are blocked. "
            "Private Endpoints are the default for every PaaS data service (Storage, "
            "Key Vault, SQL, Cosmos DB, Service Bus, Event Hubs); public network "
            "access on the resource is explicitly disabled. Private DNS zones are "
            "hub-managed and linked to all spokes to ensure consistent name "
            "resolution. ExpressRoute uses MACsec or IPsec over the private peering "
            "for sensitive workloads, and circuits are deployed in active-active "
            "across two peering locations for resiliency. DDoS Network Protection "
            "is enabled on every virtual network hosting an internet-facing endpoint."
        ),
    ),
    (
        "3. DATA PROTECTION AND ENCRYPTION",
        (
            "Data classification drives every protection control. Each workload "
            "maintains a data inventory mapping datasets to one of four sensitivity "
            "tiers (public, internal, confidential, restricted), and the tier "
            "dictates encryption, key management, retention, and egress controls. "
            "Encryption at rest is mandatory for all storage services and uses "
            "customer-managed keys (CMK) stored in Azure Key Vault Premium with "
            "HSM-backed keys for confidential and restricted tiers; platform-managed "
            "keys are acceptable only for public tier non-production workloads. "
            "Key rotation is automated on a maximum 12-month cadence with versioning "
            "preserved to support decryption of historical data. Encryption in "
            "transit requires TLS 1.2 or higher with modern cipher suites; Azure "
            "service endpoints that still permit TLS 1.0/1.1 must be explicitly "
            "overridden. Sensitive fields within databases use Always Encrypted, "
            "client-side envelope encryption, or Azure SQL column-level encryption "
            "to protect against insider DBAs. Secrets — connection strings, API "
            "keys, certificates — live exclusively in Key Vault; configuration "
            "files and ARM/Bicep templates must reference them via Key Vault "
            "references rather than embedding values. Backups, snapshots, and "
            "exported datasets inherit the source encryption posture; cross-region "
            "replication must remain within approved data-residency boundaries. "
            "Data loss prevention is enforced through Microsoft Purview policies "
            "on Microsoft 365 and via egress controls on Azure storage."
        ),
    ),
    (
        "4. STORAGE BEST PRACTICES",
        (
            "Azure Storage accounts are provisioned per workload and per environment "
            "with the minimum redundancy that satisfies the recovery objectives: "
            "ZRS for regional resilience, GZRS for cross-region durability of "
            "restricted-tier data, and LRS only for ephemeral or reproducible data. "
            "Account-level settings must enforce TLS 1.2, disable shared-key access "
            "in favour of Microsoft Entra authentication, disable public blob "
            "access, and require infrastructure encryption (double encryption) for "
            "confidential and restricted tiers. Blob soft delete is enabled with a "
            "minimum 14-day retention; container soft delete and versioning are "
            "enabled for any container hosting non-reproducible data, and immutable "
            "blob storage with time-based retention policies is used for audit and "
            "compliance datasets. Lifecycle management policies tier cool data to "
            "Cool or Archive after defined inactivity windows to control cost. "
            "Azure Files shares use identity-based authentication via Microsoft "
            "Entra Kerberos or AD DS; NTLM-only access is prohibited. Managed "
            "disks default to Premium SSD v2 for production workloads with "
            "host-based caching matched to the workload pattern, and ephemeral OS "
            "disks are used for stateless scale-set workers. Disk encryption is "
            "always on; CMK with Azure Disk Encryption Sets is required for "
            "confidential and restricted tiers. Storage accounts are reached only "
            "over Private Endpoints, with selected-networks firewall rules as a "
            "secondary defence."
        ),
    ),
    (
        "5. COMPUTE AND VM HARDENING",
        (
            "Virtual machines are the exception, not the default — workloads "
            "should be evaluated against managed compute platforms (App Service, "
            "Container Apps, Functions, AKS) before approving IaaS. Where VMs are "
            "justified, they must run from a curated, regularly updated golden "
            "image produced by Azure Image Builder or Azure Compute Gallery, with "
            "CIS Level 1 or higher hardening baked in. Trusted Launch is enabled "
            "for all Generation 2 VMs, providing Secure Boot, vTPM, and boot "
            "integrity monitoring; Confidential VMs are selected for workloads "
            "processing restricted-tier data. Just-in-Time VM access through "
            "Microsoft Defender for Cloud is mandatory; persistent inbound RDP or "
            "SSH ports on NSGs are prohibited and Azure Bastion is the only "
            "supported management path. Guest patching is performed by Azure "
            "Update Manager on a defined maintenance window — critical and "
            "security patches within 14 days of release, others within 30. "
            "Endpoint protection via Microsoft Defender for Servers Plan 2 is "
            "deployed to every VM and scale-set instance and forwards alerts to "
            "the central SOC workspace. VM extensions are constrained by Azure "
            "Policy to an approved list, and the Custom Script Extension is "
            "blocked outside of break-glass scenarios because it can execute "
            "arbitrary code. Auto-shutdown schedules are applied to all non-"
            "production VMs to control spend."
        ),
    ),
    (
        "6. CONTAINER AND KUBERNETES SECURITY",
        (
            "AKS clusters are deployed as private clusters with the API server "
            "exposed only over a Private Endpoint; public API endpoints are not "
            "permitted in production. Workload identity uses the Microsoft Entra "
            "Workload Identity integration so pods receive scoped tokens through "
            "federated credentials instead of managing service-principal secrets "
            "in cluster. Cluster authentication is integrated with Microsoft Entra "
            "ID with Azure RBAC for Kubernetes authorisation; local Kubernetes "
            "accounts are disabled. Each cluster carries Azure Policy for "
            "Kubernetes with the baseline and restricted pod-security profiles "
            "applied, denying privileged containers, host networking, host paths, "
            "and the use of the default service account. Container images are "
            "pulled exclusively from Azure Container Registry instances with "
            "Microsoft Defender for Containers scanning enabled and quarantine on "
            "vulnerable images; image signatures are verified through Notary "
            "v2/Notation policies before deployment. Network policy (Calico or "
            "Azure Network Policy) is enabled per namespace with default-deny "
            "rules and explicit egress allow lists. Cluster upgrades follow the "
            "AKS supported version policy: minor versions are upgraded within 90 "
            "days of release, and the node image is rotated monthly through "
            "auto-upgrade. Secrets are sourced through the Secret Store CSI Driver "
            "backed by Key Vault rather than Kubernetes secrets. Cluster diagnostic "
            "logs and audit logs flow to the central Log Analytics workspace."
        ),
    ),
    (
        "7. SERVERLESS AND APP SERVICE GUIDANCE",
        (
            "App Service Plans and Function premium/elastic plans live inside a "
            "regional VNet with VNet integration enabled and outbound traffic "
            "routed through the hub firewall. Inbound traffic to App Service apps "
            "is restricted by access-restriction rules combined with Private "
            "Endpoints; public access is disabled where the front door (Azure "
            "Front Door, Application Gateway, or APIM) provides ingress. Functions "
            "default to the Flex Consumption or Premium plan when VNet integration, "
            "private storage, or long-running execution is needed; the Consumption "
            "plan is reserved for stateless, public, low-sensitivity scenarios. "
            "System-assigned managed identity is enabled on every app and is the "
            "only mechanism used to call downstream Azure services. Application "
            "settings store no secrets — Key Vault references are used universally. "
            "HTTPS-only and minimum TLS 1.2 are enforced, FTP/FTPS deployment is "
            "disabled, and SCM is locked to specific source IPs or front-door "
            "headers. Deployment slots provide blue-green releases; production "
            "swaps occur only after automated smoke tests succeed. Diagnostic "
            "settings stream application logs, HTTP logs, and platform metrics to "
            "Log Analytics. Auto-heal rules and health-check endpoints are "
            "configured so the platform can recycle unhealthy instances automatically."
        ),
    ),
    (
        "8. DATABASE AND DATA PLATFORM",
        (
            "Relational workloads prefer Azure SQL Managed Instance or Azure "
            "Database for PostgreSQL Flexible Server, deployed inside a delegated "
            "subnet with Private Endpoints and no public access. Transparent Data "
            "Encryption uses customer-managed keys for confidential and restricted "
            "data; auditing and Microsoft Defender for SQL are enabled with alerts "
            "wired to the SOC. Authentication is Microsoft Entra ID only — SQL "
            "logins and built-in admin accounts are disabled after bootstrap and "
            "managed identities are used for service-to-database calls. Cosmos DB "
            "is provisioned with the minimum consistency level required (Session "
            "by default), Private Endpoint access only, and regional failover "
            "groups for tier-1 workloads; throughput is governed by autoscale "
            "with per-container caps to prevent runaway cost. Synapse and Fabric "
            "workspaces enforce managed VNet, managed Private Endpoints to "
            "linked stores, and workspace-level Microsoft Entra integration. "
            "Backup retention is defined per RPO/RTO requirement: short-term "
            "point-in-time recovery on the platform plus long-term retention "
            "policies for compliance datasets. Schema migrations flow through "
            "version-controlled pipelines with approval gates; no ad-hoc DDL is "
            "executed against production. Database connection strings are stored "
            "in Key Vault and consumed via references."
        ),
    ),
    (
        "9. OBSERVABILITY MONITORING AND LOGGING",
        (
            "Every workload emits telemetry to a central Log Analytics workspace "
            "in the same geography, with workspace-based Application Insights for "
            "application telemetry. Platform diagnostic settings are enabled on "
            "every resource via Azure Policy and forward activity logs, resource "
            "logs, and metrics; absence of diagnostic settings is an audit failure. "
            "Retention is tiered: 30–90 days hot, 1–2 years archive, with longer "
            "retention for security and compliance datasets driven by data-class "
            "policy. Microsoft Sentinel sits over the same workspace for "
            "security analytics, with connectors enabled for Microsoft Entra ID, "
            "Microsoft 365, Defender for Cloud, and the workload data planes. "
            "Service-level objectives are defined per critical user journey with "
            "SLI dashboards in Azure Workbooks or Grafana; alerts fire on burn-rate "
            "rather than raw thresholds to limit noise. Each alert is "
            "wired through Azure Monitor action groups to the team's on-call "
            "rotation in PagerDuty, ServiceNow, or Teams, with runbooks linked "
            "from the alert. Distributed tracing uses OpenTelemetry; correlation "
            "IDs traverse every service hop and surface in logs. Resource health "
            "and service health alerts inform incident response of platform-side "
            "events independent of workload telemetry."
        ),
    ),
    (
        "10. BACKUP DISASTER RECOVERY AND BUSINESS CONTINUITY",
        (
            "Every production workload declares an RPO and RTO during design "
            "review; the declared values drive the choice of redundancy, backup "
            "frequency, and DR pattern. Azure Backup protects VMs, Azure Files, "
            "SQL in VM, and PostgreSQL/MySQL Flexible Servers with vaulted backups "
            "for ransomware protection, immutability turned on, and cross-region "
            "restore enabled for tier-1 workloads. PaaS data stores use the "
            "platform's native protection (SQL PITR, Cosmos DB continuous backup, "
            "Storage soft delete + versioning + object replication) configured to "
            "match the declared RPO. Recovery Services Vaults and Backup Vaults "
            "are deployed per region with diagnostic settings forwarded to the "
            "central workspace. DR patterns range from active-active across paired "
            "regions for tier-0 workloads, to active-passive with Azure Site "
            "Recovery for tier-1 VM-based workloads, to backup-and-rebuild for "
            "tier-2 workloads where IaC enables rapid reconstruction. Every "
            "workload tests its DR plan at least annually with a documented "
            "tabletop or live failover exercise; results feed back into the SLO "
            "and architecture review. Backup credentials and recovery keys are "
            "stored in a separate Key Vault from the production keys to guard "
            "against blast radius."
        ),
    ),
    (
        "11. COST MANAGEMENT AND TAGGING",
        (
            "Cost ownership is established at design time. Every subscription, "
            "resource group, and resource carries a mandatory tag set: "
            "``costCenter``, ``owner``, ``environment``, ``workload``, "
            "``dataClassification``, and ``slaTier``. Azure Policy denies resource "
            "creation when any mandatory tag is missing, and inheritance policies "
            "propagate resource-group tags to child resources. Budgets are "
            "configured per subscription and per workload with alerts at 50, 80, "
            "and 100 percent of forecast, and action groups that notify the FinOps "
            "channel and the workload owner. Reservations and savings plans are "
            "evaluated quarterly for steady-state compute; spot VMs are used for "
            "fault-tolerant batch workloads. Idle and right-sizing recommendations "
            "from Azure Advisor are reviewed monthly and either actioned or "
            "documented as a deliberate exception. Non-production environments "
            "run on schedules: dev/test compute auto-shuts down outside business "
            "hours, and ephemeral environments are torn down by the IaC pipeline "
            "when feature branches close. Cost anomaly detection is enabled and "
            "anomalies trigger automated investigation tasks. Chargeback and "
            "showback reports flow from Microsoft Cost Management exports into "
            "the central FinOps Power BI workspace."
        ),
    ),
    (
        "12. NAMING CONVENTIONS AND RESOURCE ORGANIZATION",
        (
            "Resource names follow a deterministic pattern: ``<resourceType>-"
            "<workload>-<environment>-<region>-<instance>`` using approved "
            "abbreviations (e.g. ``stg`` for Storage, ``kv`` for Key Vault, "
            "``aks`` for Kubernetes Service), kebab-case where the resource "
            "permits it, and lowercase otherwise. Global resources whose names "
            "must be unique (Storage accounts, Key Vaults, App Services, "
            "Container Registries) append a short hash suffix derived from the "
            "subscription ID to avoid collisions. Resource groups bound the "
            "lifecycle of a related set of resources — never a mix of permanent "
            "and ephemeral resources in the same group. Subscriptions are aligned "
            "to environment and workload tier: shared services, production, "
            "non-production, and sandbox each get their own subscription scope, "
            "and management groups arrange them under a landing-zone hierarchy "
            "(platform, application, sandbox, decommissioned). Region selection "
            "is driven by data residency and latency requirements; deployments "
            "use Azure region pairs for cross-region replication. Tags supplement "
            "but do not replace the naming convention. Resource locks (CanNotDelete "
            "for production data, ReadOnly for the landing-zone backbone) prevent "
            "accidental deletion. The naming convention and abbreviation table "
            "live in the platform handbook and are enforced by an Azure Policy "
            "regex match."
        ),
    ),
    (
        "13. DEVOPS CI CD AND INFRASTRUCTURE AS CODE",
        (
            "All Azure infrastructure is defined as code — Bicep is the default "
            "for greenfield landing-zone and workload deployments; Terraform is "
            "acceptable where the team has existing investment, and ARM JSON is "
            "deprecated. Module reuse is mandatory: workloads consume versioned "
            "modules published from the platform team's private registry rather "
            "than copy-pasting resource blocks. Source of truth lives in Git "
            "with branch protection on the main branch — required reviewers, "
            "signed commits, status checks for linting (Bicep linter, "
            "tflint), unit tests, security scanning (PSRule for Azure, Checkov, "
            "tfsec), and a what-if/plan output reviewed in the PR. Pipelines run "
            "on managed agents (Azure DevOps or GitHub Actions) authenticated to "
            "Azure with workload-identity federation; long-lived service "
            "principals and PATs are not permitted. Deployments are promoted "
            "through environments (dev, test, stage, prod) with approval gates "
            "and automated smoke tests after each stage. Secrets used during "
            "deployment are pulled from Key Vault at runtime through the "
            "federated identity, never stored in pipeline variables. Production "
            "deployments occur during defined change windows and emit "
            "deployment annotations to Application Insights so observability "
            "dashboards can correlate behaviour changes."
        ),
    ),
    (
        "14. COMPLIANCE GOVERNANCE AND POLICY",
        (
            "Azure Policy is the primary governance mechanism. The platform "
            "team assigns regulatory and corporate initiatives (ISO 27001, SOC 2, "
            "PCI DSS where applicable, and the internal landing-zone baseline) "
            "at the management-group root with deny and audit effects calibrated "
            "to the environment tier — deny in production, audit in sandbox. "
            "Policy assignments expose parameters that workload teams set per "
            "subscription rather than rewriting policies. Microsoft Defender for "
            "Cloud regulatory compliance dashboards are reviewed monthly, with "
            "non-compliant resources triaged within an SLA tied to severity. "
            "Microsoft Purview catalogues data assets and lineage, applying "
            "sensitivity labels that drive downstream access and protection. "
            "Audit logs (activity log, Microsoft Entra sign-in and audit, "
            "Defender alerts) are forwarded to an immutable archive with a "
            "minimum 7-year retention for SOX-relevant subscriptions. Exception "
            "management is formal: every policy exemption carries a justification, "
            "owner, compensating control, and expiry date, and exemptions are "
            "reviewed quarterly. Internal and external audits draw from the same "
            "evidence store. Cost of compliance is itself reported so the "
            "business can weigh control investments."
        ),
    ),
    (
        "15. AI AND AGENT WORKLOADS GOVERNANCE",
        (
            "Generative AI and agent workloads built on Azure AI Foundry, Azure "
            "OpenAI, or Azure Machine Learning inherit every control in this "
            "document and add domain-specific governance. Foundry projects and "
            "Azure OpenAI resources are deployed regionally to satisfy data "
            "residency, with Private Endpoint access and Microsoft Entra "
            "authentication; API keys are disabled in favour of managed-identity "
            "access from caller workloads. Model selection prefers regional GA "
            "deployments — at the time of writing Canada Central does not list "
            "``gpt-5.4-pro``, so production workloads default to "
            "``gpt-5.3-chat-1`` (or the equivalent successor already deployed in "
            "``foundry-cc-canada``) and document the fallback in a prompt "
            "contract; deployments auto-upgrade when the target model becomes "
            "available. Content safety filters run on every request and response; "
            "prompt-injection defences and jailbreak detection are enabled, and "
            "violations are logged to the security workspace. Grounding data "
            "retrieved through Azure AI Search or other RAG sources respects the "
            "same data classification: restricted-tier corpora live in dedicated "
            "indexes with filterable security trimming. Agent and tool "
            "definitions are version-controlled prompt contracts that declare "
            "intent, inputs, outputs, edge cases, and acceptance criteria; "
            "changes go through code review. Evaluation pipelines run automated "
            "groundedness, relevance, and safety scoring before promotion. "
            "Cost and quota are monitored per deployment with budgets and rate "
            "limits aligned to business value, and PII redaction is applied to "
            "telemetry to satisfy privacy commitments."
        ),
    ),
]


def build() -> Path:
    doc = Document()

    # Title and short preamble — neither line is UPPERCASE-numeric so the
    # extractor will ignore them.
    title = doc.add_paragraph("Default Azure Infrastructure Policies")
    title.style = doc.styles["Title"]
    intro = doc.add_paragraph(
        "Originally written guidance for ARB validation. Section headers are "
        "uppercase and numbered to match the policy extractor."
    )
    intro.style = doc.styles["Normal"]

    for header, body in SECTIONS:
        h = doc.add_paragraph(header)
        h.style = doc.styles["Heading 1"]
        for run in h.runs:
            run.font.size = Pt(14)
        # Split long bodies into 2–3 paragraphs at sentence boundaries so the
        # document reads naturally and stays comfortably under 50 pages.
        sentences = [s.strip() for s in body.split('. ') if s.strip()]
        chunk_size = max(1, (len(sentences) + 2) // 3)
        for i in range(0, len(sentences), chunk_size):
            chunk = '. '.join(sentences[i:i + chunk_size])
            if not chunk.endswith('.'):
                chunk += '.'
            doc.add_paragraph(chunk)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    p = build()
    print(f"Wrote {p}")
    sys.exit(0)
